from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "_backend_doc_assets"
OUTPUT = OUT_DIR / "OptiSched_Backend_Teknik_Dokumani.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "667085"
INK = "182230"
WHITE = "FFFFFF"
GREEN = "217346"
GOLD = "9A6700"
RED = "9B1C1C"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            set_cell_margins(row.cells[index])
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def style_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_text(doc, text: str, *, bold_prefix: str | None = None, after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, color=color)
    else:
        style_run(p.add_run(text), color=color)
    return p


def add_bullet(doc, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))


def add_callout(doc, title: str, body: str, fill=LIGHT_BLUE, title_color=NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(title), size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    style_run(p2.add_run(body), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(header), size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            style_run(p.add_run(str(value)), size=font_size)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for line_no, line in enumerate(code.splitlines()):
        if line_no:
            p.add_run().add_break()
        style_run(p.add_run(line), size=8.5, font="Courier New", color="24292F")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def font(size: int, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, fill, size=30, bold=False):
    fnt = font(size, bold)
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        text,
        font=fnt,
        fill=fill,
        spacing=8,
        align="center",
    )


def rounded_box(draw, box, fill, outline=BLUE, radius=22, width=4):
    if isinstance(outline, str) and not outline.startswith("#"):
        outline = f"#{outline}"
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=BLUE, width=7):
    if isinstance(color, str) and not color.startswith("#"):
        color = f"#{color}"
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 10), (x - 18, y + 10)], fill=color)


def architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 850), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    draw.text((65, 40), "OptiSched Backend Mimarisi", font=font(48, True), fill=f"#{NAVY}")
    boxes = {
        "front": (70, 190, 390, 420),
        "api": (500, 130, 890, 480),
        "db": (1030, 100, 1510, 290),
        "solver": (1030, 350, 1510, 540),
        "excel": (1030, 600, 1510, 790),
    }
    rounded_box(draw, boxes["front"], f"#{LIGHT_BLUE}")
    rounded_box(draw, boxes["api"], "#DCEEFF")
    rounded_box(draw, boxes["db"], "#E8F5EC", outline=GREEN)
    rounded_box(draw, boxes["solver"], "#FFF5D9", outline=GOLD)
    rounded_box(draw, boxes["excel"], "#FDECEC", outline=RED)
    draw_centered(draw, boxes["front"], "React + Vite\nFrontend", f"#{NAVY}", 34, True)
    draw_centered(draw, boxes["api"], "FastAPI\nAPI + Adapter +\nRepository Katmanı", f"#{NAVY}", 32, True)
    draw_centered(draw, boxes["db"], "PostgreSQL\nKısıtlar, kullanıcılar,\nrezervasyonlar, programlar", f"#{GREEN}", 28, True)
    draw_centered(draw, boxes["solver"], "OR-Tools CP-SAT\nalgorithm/shared/\noptisched_core", f"#{GOLD}", 29, True)
    draw_centered(draw, boxes["excel"], "Dönem Excel Dosyaları\nORTAK + 12 Bölüm", f"#{RED}", 29, True)
    arrow(draw, (390, 305), (500, 305))
    arrow(draw, (890, 210), (1030, 195))
    arrow(draw, (890, 340), (1030, 440))
    arrow(draw, (890, 430), (1030, 695))
    image.save(path)


def flow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 930), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    draw.text((65, 35), "Tek Tıklamada İki Aşamalı Program Üretimi", font=font(46, True), fill=f"#{NAVY}")
    nodes = [
        ((70, 165, 350, 345), "1. Dönem seçimi\nGüz / Bahar", LIGHT_BLUE, BLUE),
        ((440, 165, 760, 345), "2. ORTAK Excel\nsatırlarını al", "FDECEC", RED),
        ((850, 165, 1170, 345), "3. Ortak dersleri\nCP-SAT ile çöz", "FFF5D9", GOLD),
        ((1260, 165, 1530, 345), "4. Ortak öğretim\nüyesi doluluğunu kilitle", "E8F5EC", GREEN),
        ((215, 530, 560, 735), "5. 12 bölüm Excel'ini\nsırayla oku ve çöz", LIGHT_BLUE, BLUE),
        ((650, 530, 995, 735), "6. Bölümler arası\nöğretim üyesi çakışmasını\nengelle", "FFF5D9", GOLD),
        ((1085, 530, 1430, 735), "7. Doğrula, JSONB'ye\nkaydet ve frontend'e dön", "E8F5EC", GREEN),
    ]
    for box, text, fill, outline in nodes:
        rounded_box(draw, box, f"#{fill}", outline=outline)
        draw_centered(draw, box, text, f"#{outline}", 27, True)
    for start, end in [
        ((350, 255), (440, 255)),
        ((760, 255), (850, 255)),
        ((1170, 255), (1260, 255)),
        ((1395, 345), (1395, 440)),
        ((1395, 440), (390, 440)),
        ((390, 440), (390, 530)),
        ((560, 632), (650, 632)),
        ((995, 632), (1085, 632)),
    ]:
        arrow(draw, start, end)
    image.save(path)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(p.add_run("OPTISCHED  |  BACKEND TEKNİK REFERANSI"), size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(p.add_run("Mehmet Akif Duran  |  Haziran 2026  |  "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "architecture.png"
    flow = ASSET_DIR / "schedule_flow.png"
    architecture_diagram(architecture)
    flow_diagram(flow)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)

    doc.add_paragraph().paragraph_format.space_after = Pt(86)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    style_run(p.add_run("BİTİRME PROJESİ TEKNİK DOKÜMANI"), size=11, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run("OptiSched Backend"), size=31, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    style_run(p.add_run("Mimari, API, PostgreSQL, CP-SAT Algoritması ve İş Akışları"), size=15, color=DARK_BLUE)
    add_callout(
        doc,
        "Dokümanın amacı",
        "Bu belge, projenin backend tarafını sunum ve jüri hazırlığı için uçtan uca açıklar. "
        "Kodun mevcut davranışını, veri akışını, kısıt modelini, doğrulama mekanizmasını ve bilinen teknik sınırları birbirinden ayırır.",
        fill="F4F6F9",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(105)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Hazırlayan: Mehmet Akif Duran\nHaziran 2026"), size=11, bold=True, color=MUTED)

    page_break(doc)
    doc.add_heading("1. Yönetici Özeti", level=1)
    add_callout(
        doc,
        "30 saniyelik anlatım",
        "OptiSched, üniversite ders programlarını PostgreSQL'de tutulan kurallara göre üreten bir FastAPI servisidir. "
        "OR-Tools CP-SAT çözücüsü önce ortak dersleri, ardından 12 mühendislik bölümünü sırayla planlar. "
        "Öğretim üyesi, dönem/şube, süre, bölünmüş ders ve zaman kısıtları çözüm sırasında uygulanır; sonuçlar bağımsız doğrulayıcıdan geçirilerek JSONB program belgeleri halinde saklanır.",
    )
    add_text(doc, "Backendin dört ana sorumluluğu vardır:")
    for item in [
        "Frontend'e kimlik doğrulama, ders, program, kısıt ve rezervasyon API'leri sunmak.",
        "Düzenlenebilir algoritma kısıtlarını PostgreSQL'den okumak ve yine PostgreSQL'e kaydetmek.",
        "Excel ve veritabanı verilerini OR-Tools çözücüsünün beklediği tablo biçimine dönüştürmek.",
        "Üretilen programları doğrulamak, frontend modeline çevirmek ve dönem/bölüm bazında kalıcılaştırmak.",
    ]:
        add_bullet(doc, item)
    add_table(
        doc,
        ["Katman", "Teknoloji", "Temel rol"],
        [
            ["HTTP/API", "FastAPI 0.115", "Endpoint, CORS, hata yanıtı ve bağımlılık yönetimi"],
            ["Şema", "Pydantic 2", "İstek/yanıt doğrulama ve tip sözleşmesi"],
            ["Veri", "PostgreSQL + psycopg2", "Kullanıcı, ders, kısıt, rezervasyon ve program saklama"],
            ["Algoritma", "OR-Tools CP-SAT 9.12", "Uygun ders-zaman yerleşimi ve amaç fonksiyonu"],
            ["Dönüşüm", "pandas + openpyxl", "Excel okuma, veri normalizasyonu ve çıktı işleme"],
        ],
        [1.1, 1.7, 3.7],
    )

    doc.add_heading("2. Sistem Mimarisi", level=1)
    doc.add_picture(str(architecture), width=Inches(6.45))
    p = doc.add_paragraph("Şekil 1. OptiSched backend bileşenleri ve temel veri yönleri.")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(), size=8.5, color=MUTED)
    add_text(
        doc,
        "FastAPI, frontend ile veri kaynakları arasında orkestrasyon katmanıdır. "
        "Algoritma doğrudan HTTP ya da PostgreSQL bilmez; pandas DataFrame ve yapılandırma nesneleri üzerinden çalışır. "
        "Bu ayrım, çözücünün terminal uygulamasında ve web backendinde ortak kullanılmasını sağlar.",
    )
    doc.add_heading("2.1 Kaynak Kod Yerleşimi", level=2)
    add_table(
        doc,
        ["Konum", "Sorumluluk"],
        [
            ["backend/app/main.py", "FastAPI uygulaması ve bütün HTTP endpointleri"],
            ["backend/app/models.py", "Pydantic istek/yanıt modelleri"],
            ["backend/app/auth.py", "Parola doğrulama ve bearer token üretimi"],
            ["backend/app/*_repository.py", "PostgreSQL sorguları ve kalıcılık"],
            ["backend/app/scheduler_adapter.py", "DB/Excel verisini algoritmaya bağlayan adaptör"],
            ["algorithm/shared/optisched_core/", "CP-SAT model, doğrulama, blok sıkıştırma ve sınıf atama yardımcıları"],
            ["database/migrations/", "Kısıt, program ve rezervasyon tablolarını ekleyen migrationlar"],
            ["real_deal_database/", "Güz/Bahar bölüm Excel dosyaları"],
        ],
        [2.45, 4.05],
        9.2,
    )

    page_break(doc)
    doc.add_heading("3. Uçtan Uca Program Üretim Akışı", level=1)
    doc.add_picture(str(flow), width=Inches(6.45))
    p = doc.add_paragraph("Şekil 2. Kullanıcı açısından tek işlem, backend açısından iki aşamalı fakülte çözümü.")
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(), size=8.5, color=MUTED)
    doc.add_heading("3.1 Aşama 1: Ortak Dersler", level=2)
    for item in [
        "Frontend, seçilen dönemin ORTAK Excel satırlarını `common_rows` alanıyla gönderir.",
        "Backend bu satırları teori/laboratuvar saatlerine göre ders parçalarına dönüştürür.",
        "Ortak ders çözümünde dönem-şube çakışması kapatılır; öğretim üyesi, zaman ve bağlantılı ders kuralları uygulanır.",
        "Ortak ders çözümü başarısızsa bölüm aşamasına geçilmez ve açıklayıcı hata döndürülür.",
    ]:
        add_number(doc, item)
    doc.add_heading("3.2 Aşama 2: Bölümler", level=2)
    for item in [
        "Seçilen dönemin 12 bölüm Excel dosyası sırayla okunur: BİL, EEM, END, CE, BME, MAK, AI, CSE, BENG, EEE, IND ve ME.",
        "Ortak derslerde dolu olan öğretim üyesi zamanları bölüm çözümlerine yasak zaman olarak aktarılır.",
        "Her bölüm çözüldükçe yeni öğretim üyesi dolulukları sonraki bölümlere taşınır.",
        "Son aşamada birleşik program üzerinde bölümler arası öğretim üyesi çakışması tekrar kontrol edilir.",
        "Ortak ve bölüm programları `published_schedule_documents` tablosuna ayrı JSONB belgeleri olarak kaydedilir.",
    ]:
        add_number(doc, item)
    add_callout(
        doc,
        "Dönem seçiminin etkisi",
        "`fall` yalnızca tek dönem numaralı; `spring` yalnızca çift dönem numaralı dersleri tekil çalıştırma akışında filtreler. "
        "Fakülte akışında ise doğru `GUZ` veya `BAHAR` Excel dizini seçilir. Akademik yıl görünürde kaldırılmış olsa da backend program anahtarını şu anda `0000-0000|fall/spring` biçiminde üretmektedir.",
        fill="FFF8E8",
        title_color=GOLD,
    )

    doc.add_heading("4. API Sözleşmesi", level=1)
    endpoint_rows = [
        ["GET", "/api/health", "Hayır", "Servis sağlık kontrolü"],
        ["POST", "/api/auth/login", "Hayır", "Kullanıcı doğrular, token döndürür"],
        ["GET", "/api/auth/me", "Evet", "Aktif hesabı döndürür"],
        ["POST", "/api/auth/logout", "Hayır", "204 döndürür; token sunucuda tutulmaz"],
        ["GET", "/api/courses", "Hayır", "Ders ve bölüm eşleşmelerini listeler"],
        ["GET", "/api/rooms", "Hayır", "Derslikleri kapasite ve türle listeler"],
        ["GET", "/api/rooms/availability", "Hayır", "Tarih/saat için uygunluk hesaplar"],
        ["GET", "/api/reservations", "Evet", "Rezervasyonları listeler"],
        ["POST", "/api/reservations", "Evet", "Çakışma kontrollü rezervasyon oluşturur"],
        ["DELETE", "/api/reservations/{id}", "Evet", "Yetkiye göre rezervasyon siler"],
        ["GET", "/api/scheduler/inputs", "Hayır", "DB'den algoritma ders girdilerini üretir"],
        ["POST", "/api/scheduler/inputs", "Evet", "Tek ders girdisi ekler/günceller"],
        ["GET/PUT", "/api/scheduler/common", "PUT: Evet", "Ortak programı okur/yazar"],
        ["GET/PUT", "/api/scheduler/dept", "PUT: Evet", "Bölüm programlarını okur/yazar"],
        ["GET/PUT", "/api/scheduler/linked", "PUT: Evet", "Bağlantılı ders gruplarını yönetir"],
        ["GET/PUT", "/api/scheduler/constraint-bundle", "Hayır", "Kısıt paketini okur/yazar"],
        ["POST", "/api/scheduler/run", "Hayır", "Tek kapsamlı çözücü akışını çalıştırır"],
        ["POST", "/api/scheduler/run-university", "Opsiyonel", "Ortak + 12 bölüm akışını çalıştırır ve kaydeder"],
    ]
    add_table(doc, ["Metot", "Endpoint", "Kimlik", "Amaç"], endpoint_rows, [0.65, 2.55, 0.8, 2.5], 8.1)
    add_callout(
        doc,
        "Sunumda vurgulanacak nokta",
        "Frontend'in algoritma kurallarını doğrudan çözücüye göndermesi yerine, kısıt paketi önce PostgreSQL'e yazılır; backend çalıştırma anında paketi yeniden veritabanından okur. Böylece kullanılan kural seti kayıt altındadır.",
    )

    page_break(doc)
    doc.add_heading("5. Veri Modelleri ve Dönüşümler", level=1)
    add_table(
        doc,
        ["Model", "Önemli alanlar", "Kullanım"],
        [
            ["AlgorithmInput", "kod, ad, haftalık saat, dönem, şube, öğretim üyesi, tür", "DB ders bilgisinin çözücü girdisine taşınması"],
            ["ConstraintBundle", "settings, systemRules, constraints, linkedGroups, splits", "Bir dönem ve bölüm için bütün algoritma ayarları"],
            ["ScheduledCourse", "gün, başlangıç/bitiş, şube, sınıf, bölüm, dinleyici bölümler", "Frontend'in gösterdiği program satırı"],
            ["UniversitySchedulerRunRequest", "term, courses, common_rows, view_department, common_only", "Fakülte çapında çalıştırma isteği"],
            ["ScheduleStats", "süre, yerleşen görev, hata, uyarı, öğretim üyesi saatleri", "Çalıştırma özeti ve gözlemlenebilirlik"],
            ["ReservationWrite", "derslik, tarih, saatler, ders, öğretim üyesi", "Rezervasyon oluşturma isteği"],
        ],
        [1.55, 2.65, 2.3],
        8.8,
    )
    doc.add_heading("5.1 Excel Normalizasyonu", level=2)
    for item in [
        "`Ders Kodu` büyük harfe ve boşluksuz biçime getirilir.",
        "`Şube` sayısallaştırılır; eksikse 1 kabul edilir.",
        "`t_hour` ve `l_hour` toplamı haftalık süreyi belirler; bunlar yoksa `Kredi` kullanılır.",
        "4 saat 2+2, 5 saat 2+3, 6 saat 3+3 olarak bölünebilir; DB'deki özel split tercihi önceliklidir.",
        "Öğrenci sayısında önce `Sınıf Mevcudu`, sonra `Kontenjan`, son olarak 0 kullanılır.",
        "Bitirme Projesi I/II ve Graduation Project I/II adları girdi, çıktı ve kayıt aşamalarında filtrelenir.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("5.2 Frontend Çıktısına Dönüşüm", level=2)
    add_text(
        doc,
        "Çözücü her ders saati için ayrı satır üretir. `compress_schedule_blocks` ardışık saatleri tek ders bloğunda birleştirir. "
        "Adaptör gün ve saat indekslerini `Mon`/`09:00` biçimine çevirir, öğretim üyesi yoksa `Atanmamış`, sınıf yoksa varsayılan `TBA` değerini kullanır.",
    )

    page_break(doc)
    doc.add_heading("6. PostgreSQL Tasarımı", level=1)
    add_text(
        doc,
        "Mevcut üniversite şeması dersler, şubeler, öğretim üyeleri, bölümler, kullanıcılar ve derslikler için ana kaynaktır. "
        "Proje migrationları bu şemayı silmeden yalnızca çizelgeleme ve rezervasyon tablolarını ekler.",
    )
    add_table(
        doc,
        ["Tablo", "Görev", "Önemli ilişki"],
        [
            ["academic_periods", "Akademik yıl + güz/bahar anahtarı", "year_label + term benzersiz"],
            ["scheduler_constraint_sets", "Takvim, süre, ağırlık ve hariç prefix ayarları", "period + department benzersiz"],
            ["scheduler_system_rules", "Açılıp kapanabilen kural bayrakları", "constraint_set'e bağlı"],
            ["scheduler_time_constraints", "Hoca müsait değil / ders sabit / ders yasak", "hedef + gün + saat"],
            ["scheduler_course_split_preferences", "Dersin haftalık saat bölme tercihi", "ders kodu + sıra"],
            ["linked_course_groups", "Birlikte başlayan ders grubu üst kaydı", "döneme bağlı UUID"],
            ["linked_group_courses", "Gruba bağlı ders kodları", "group_id + course_code"],
            ["published_schedule_documents", "Ortak ve bölüm programlarını JSONB saklar", "period + department benzersiz"],
            ["classroom_reservations", "Tarihli derslik rezervasyonları", "derslik, tarih, text[] saat"],
        ],
        [2.15, 2.75, 1.6],
        8.5,
    )
    add_callout(
        doc,
        "`HAVUZ` kapsamı",
        "Ortak dersler veritabanında özel bir bölüm kodu olan `HAVUZ` ile temsil edilir. Diğer programlar gerçek bölüm kodlarıyla saklanır.",
        fill="F4F6F9",
    )

    page_break(doc)
    doc.add_heading("7. Kısıt Modeli ve CP-SAT", level=1)
    add_text(
        doc,
        "Her ders parçası, gün ve saat için bir Boolean karar değişkeniyle modellenir. "
        "Başlangıç değişkenleri, dersin tek günde ardışık saatlere yerleşmesini sağlar. "
        "OR-Tools CP-SAT uygun bir çözüm ararken zorunlu kuralları ihlal edemez; yumuşak hedefleri ise ağırlıklı maliyet olarak minimize eder.",
    )
    doc.add_heading("7.1 Zorunlu Kurallar", level=2)
    hard_rows = [
        ["H1", "Haftalık saat", "Her ders parçası tam olarak beklenen saat kadar yerleşir."],
        ["H2", "Tek gün + ardışıklık", "Bir parça aynı gün içinde kesintisiz saatlere yerleşir."],
        ["H3", "Bölünmüş parçalar", "Aynı dersin 2+2 gibi parçaları farklı günlerde olur."],
        ["H4", "Öğretim üyesi", "Aynı öğretim üyesi aynı anda iki farklı ders veremez."],
        ["H5", "Zorunlu ders", "Aynı dönem/şubenin zorunlu dersleri çakışamaz."],
        ["H6", "Seçmeli-zorunlu", "Aynı dönem/şubenin seçmelisi zorunlu dersle çakışamaz."],
        ["H7", "Kilitli yerleşim", "Önceden sabitlenen ders saati korunur."],
        ["H8", "Derslik", "Atama yapılmışsa aynı derslik aynı anda iki derse verilemez."],
    ]
    add_table(doc, ["Kod", "Kural", "Açıklama"], hard_rows, [0.55, 1.5, 4.45], 9)
    doc.add_heading("7.2 Veritabanından Gelen Dinamik Kurallar", level=2)
    dynamic_rows = [
        ["instructor_unavailable", "Belirli öğretim üyesi için gün/saat yasaklanır."],
        ["course_fixed", "Belirli ders gün/saatte zorunlu tutulur."],
        ["course_blocked", "Belirli dersin gün/saatte yerleşmesi engellenir."],
        ["linkedGroups", "Gruba bağlı derslerin başlangıçları eşitlenir."],
        ["splits", "Dersin toplam saati için özel bölme deseni seçilir."],
        ["systemRules", "Hoca, dönem, ardışık saat ve haftalık saat kontrolleri açılıp kapanır."],
    ]
    add_table(doc, ["Alan", "Çözücü etkisi"], dynamic_rows, [2.0, 4.5], 9)
    page_break(doc)
    doc.add_heading("7.3 Yumuşak Hedefler", level=2)
    add_table(
        doc,
        ["Hedef", "Varsayılan ağırlık", "Amaç"],
        [
            ["Yakın dönem çakışması", "8", "1-3, 2-4 gibi dönemlerin zorunlu ders çakışmasını azaltmak"],
            ["Öğrenci boşluğu", "0", "Gün içindeki ders aralarını azaltmak; maliyetli olduğu için varsayılan kapalı"],
            ["Öğle saati", "2", "12:00-13:00 kullanımını azaltmak"],
            ["Hoca yük dengesi", "1", "Öğretim saatlerini günlere daha dengeli yaymak"],
        ],
        [2.1, 1.25, 3.15],
        9,
    )
    add_callout(
        doc,
        "Neden CP-SAT?",
        "Problem; çakışmama, ardışıklık, sabitleme ve birlikte başlama gibi çok sayıda ayrık kararı aynı anda içerir. "
        "CP-SAT bu tür Boolean/tamsayılı kısıt problemleri için uygundur ve süre sınırı içinde optimal ya da uygulanabilir çözüm döndürebilir.",
    )

    doc.add_heading("8. Bağımsız Doğrulama Katmanı", level=1)
    add_text(
        doc,
        "Çözücünün sonuç üretmesi tek başına yeterli kabul edilmez. `validator.py`, CP-SAT modeline bakmadan somut gün/saat yerleşimlerini H1-H8 kurallarına göre yeniden kontrol eder. "
        "Bu yaklaşım modelleme veya son işleme hatalarını yakalamak için ikinci bir güvenlik katmanıdır.",
    )
    for item in [
        "Tek bölüm doğrulaması: saat, ardışıklık, split, hoca, dönem/şube ve kilit kontrolü.",
        "Fakülte doğrulaması: bütün bölüm çıktıları birleştirilerek öğretim üyesi çakışması aranır.",
        "Uyarılar `ScheduleStats.warnings` alanıyla frontend'e döner.",
        "Çözücü boş sonuç döndürürse `solver-infeasible`, ortak aşama boşsa `common-infeasible` hatası üretilir.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("9. Kimlik Doğrulama ve Yetki", level=1)
    add_text(
        doc,
        "Kullanıcı adıyla bulunan veritabanı kaydındaki `password_hash`, girilen parolanın SHA-256 özetiyle sabit zamanlı karşılaştırılır. "
        "Başarılı girişte kullanıcı kimliği ve son kullanma zamanı içeren Base64 payload, `APP_SECRET` ile HMAC-SHA256 imzalanır.",
    )
    add_table(
        doc,
        ["Konu", "Mevcut uygulama", "Değerlendirme"],
        [
            ["Token", "Özel HMAC imzalı bearer token", "JWT kütüphanesi kullanılmıyor; sunucu tarafı oturum kaydı yok"],
            ["Süre", "Varsayılan 28.800 saniye (8 saat)", "`TOKEN_TTL_SECONDS` ile değiştirilebilir"],
            ["Parola", "SHA-256 hash karşılaştırması", "Üretimde Argon2/bcrypt + salt tercih edilmelidir"],
            ["Çıkış", "204 No Content", "Stateless token iptal listesi olmadığı için istemci tokenı siler"],
            ["Roller", "admin, dept_chair, coordinator, instructor, secretary, viewer", "Silme işleminde rol kontrolü var; tüm endpointlerde tam RBAC yok"],
        ],
        [1.0, 2.7, 2.8],
        8.8,
    )
    add_callout(
        doc,
        "Üretim güvenliği",
        "Varsayılan geliştirme secret'ı kullanılmamalı; uzun ve rastgele `APP_SECRET` tanımlanmalı. "
        "Constraint bundle ve scheduler çalıştırma endpointlerinin bir kısmı şu an kimlik doğrulama istemediği için üretim öncesinde rol bazlı koruma eklenmelidir.",
        fill="FDECEC",
        title_color=RED,
    )

    doc.add_heading("10. Rezervasyon Sistemi", level=1)
    for item in [
        "Derslikler kapasite ve tür bilgisiyle filtrelenir.",
        "Seçilen tarih ve saatler, hem manuel rezervasyonlarla hem yayınlanmış ders programıyla karşılaştırılır.",
        "PostgreSQL `text[]` çakışma operatörü (`&&`) aynı saatlerden en az birinin dolu olup olmadığını kontrol eder.",
        "Aynı derslik+tarih için transaction seviyesinde advisory lock alınır; eşzamanlı çift rezervasyon yarışı engellenir.",
        "Çakışmada HTTP 409, bulunmayan derslikte HTTP 404 döner.",
        "Admin ve bölüm başkanı her rezervasyonu; diğer roller yalnızca kendi rezervasyonunu silebilir.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Dönem eşleme ayrıntısı",
        "Rezervasyon uygunluğu, tarihin ayına göre `spring` veya `fall` seçer ve yayınlanmış programı `0000-0000` yıl etiketi altında arar. "
        "Bu geçici yıl anahtarı ileride gerçek akademik yıl modeliyle değiştirilmelidir.",
        fill="FFF8E8",
        title_color=GOLD,
    )

    doc.add_heading("11. Sınıf Atama Durumu", level=1)
    add_text(
        doc,
        "`RoomAllocator`, kapasitesi yeterli en küçük uygun dersliği seçebilen ve aynı zaman diliminde derslik çakışmasını engelleyen bir bileşendir. "
        "Ancak mevcut web backend akışı bu bileşeni çağırmamaktadır. Sonuç DataFrame'inde `Sinif` alanı oluşmadığı için frontend modelinde oda varsayılan olarak `TBA` kalabilir.",
    )
    add_table(
        doc,
        ["Durum", "Sonuç"],
        [
            ["Algoritma sınıf ataması çağrılmıyor", "Program zamanı üretilir; fiziksel derslik garanti edilmez."],
            ["`TBA`", "To Be Announced: sınıf henüz belirlenmedi anlamındadır."],
            ["Rezervasyon kontrolü", "Yalnızca program satırında gerçek oda varsa doluluğu görebilir; `TBA` satırları yok sayılır."],
            ["Tam entegrasyon için", "Derslik DataFrame'i DB'den okunmalı, blok sıkıştırmadan sonra `RoomAllocator.allocate` çağrılmalı ve H8 doğrulanmalıdır."],
        ],
        [2.0, 4.5],
        9,
    )

    page_break(doc)
    doc.add_heading("12. Hata Yönetimi ve Kalıcılık", level=1)
    add_table(
        doc,
        ["Katman", "Davranış"],
        [
            ["DB bağlantısı", "Context manager başarılı işlemde commit, hatada rollback, her durumda close yapar."],
            ["DB genel hatası", "FastAPI exception handler HTTP 503 ve hata mesajı döndürür."],
            ["Model doğrulama", "Pydantic geçersiz dönem, saat, rol ve alanları 422 ile reddeder."],
            ["Program kaydı", "Period+bölüm çakışmasında JSONB belge upsert edilir; updated_at yenilenir."],
            ["Bitirme projesi filtresi", "Hem kaydetmeden önce hem okurken tekrar uygulanır."],
            ["Constraint kaydı", "Paket bütün olarak değiştirilir; alt zaman/split/rule kayıtları transaction içinde yeniden yazılır."],
        ],
        [1.45, 5.05],
        9.2,
    )
    doc.add_heading("12.1 Migration Stratejisi", level=2)
    add_text(
        doc,
        "Migrationlar eklemeli ve mümkün olduğunca idempotent tasarlanmıştır. `003_scheduler_schema_compatibility.sql`, aynı adlı eski tablolar farklı kolona sahipse onları `_legacy` adına taşıyıp güncel şemayı kurar. "
        "Bu, eski veriyi silmeden isim çakışmasını çözmeyi amaçlar.",
    )

    doc.add_heading("13. Testler ve Mevcut Doğrulama", level=1)
    add_table(
        doc,
        ["Test grubu", "Kapsam"],
        [
            ["test_scheduler_adapter.py", "Bitirme projesi filtresi, yüklenen ortak satırlar, aşırı zorunlu yük normalizasyonu"],
            ["test_solver_constraints.py", "Hoca müsait değil, sabit ders, bağlantılı ders, ortak ders öğrenci doluluğu"],
            ["test_constraint_repository.py", "Kısıt paketinin PostgreSQL'e yazılıp geri okunması; `RUN_DB_TESTS=1` ile"],
        ],
        [2.15, 4.35],
        9,
    )
    add_callout(
        doc,
        "Son bilinen uçtan uca doğrulama",
        "Geliştirme sırasında fakülte akışı 137 ortak satır ve 12 bölümle çalıştırılmış; 649 program oturumu, 0 bitirme projesi ve 0 uyarı ürettiği gözlemlenmiştir. "
        "Bu sonuç çalışma anına aittir; veri veya kısıt değiştiğinde yeniden test edilmelidir.",
        fill="E8F5EC",
        title_color=GREEN,
    )
    add_text(
        doc,
        "Bu doküman hazırlanırken birim testleri yeniden çalıştırılmıştır: 7 test başarılıdır. "
        "PostgreSQL round-trip entegrasyon testi, `RUN_DB_TESTS=1` verilmediği için atlanmıştır.",
        color=GREEN,
    )

    doc.add_heading("14. Bilinen Sınırlar ve Teknik Borçlar", level=1)
    risks = [
        ["Yüksek", "Sınıf atama bağlı değil", "`RoomAllocator` fakülte/web akışına entegre edilmeli; H8 gerçek çıktı üzerinde çalıştırılmalı."],
        ["Yüksek", "Bazı yazma/çalıştırma endpointleri korumasız", "Constraint ve scheduler endpointlerine rol bazlı yetki eklenmeli."],
        ["Orta", "Bölüm verisi sunucu Excel'inden okunuyor", "Frontend'den yüklenen tüm bölüm dosyaları backend'e aktarılmalı veya DB'ye import edilmelidir."],
        ["Orta", "Geçici akademik yıl anahtarı", "`0000-0000` yerine kullanıcı tarafından seçilen gerçek akademik dönem kullanılmalı."],
        ["Orta", "SHA-256 parola saklama", "Argon2id veya bcrypt, benzersiz salt ve parola yükseltme planı uygulanmalı."],
        ["Orta", "Logout token iptal etmiyor", "Kısa süreli access token + refresh token/denylist tasarımı değerlendirilmeli."],
        ["Düşük", "Test kapsamı sınırlı", "API integration, rezervasyon concurrency ve tüm fakülte regresyon testleri eklenmeli."],
        ["Düşük", "503 yanıtında ham DB mesajı", "İstemciye genel hata, log sistemine ayrıntı gönderilmeli."],
    ]
    add_table(doc, ["Öncelik", "Konu", "Öneri"], risks, [0.65, 1.9, 3.95], 8.5)

    page_break(doc)
    doc.add_heading("15. Kurulum ve Çalıştırma", level=1)
    doc.add_heading("15.1 Ortam Değişkenleri", level=2)
    add_table(
        doc,
        ["Değişken", "Örnek", "Amaç"],
        [
            ["DATABASE_URL", "postgresql://user:pass@localhost:5432/optisched", "PostgreSQL bağlantısı"],
            ["FRONTEND_ORIGIN", "http://127.0.0.1:5173", "CORS izinli frontend"],
            ["APP_SECRET", "uzun-rastgele-secret", "Token HMAC imzası"],
            ["TOKEN_TTL_SECONDS", "28800", "Token ömrü"],
        ],
        [1.55, 3.0, 1.95],
        8.8,
    )
    doc.add_heading("15.2 Backend", level=2)
    add_code(
        doc,
        """cd /Users/akifduran/Desktop/new_project
python3 -m venv .venv
source .venv/bin/activate
pip install -r backend/requirements.txt

export DATABASE_URL='postgresql://USER:PASSWORD@localhost:5432/optisched'
for migration in database/migrations/*.sql; do
  psql "$DATABASE_URL" -v ON_ERROR_STOP=1 -f "$migration"
done

PYTHONPATH=. uvicorn backend.app.main:app --reload --host 127.0.0.1 --port 8000""",
    )
    doc.add_heading("15.3 Frontend Bağlantısı", level=2)
    add_code(
        doc,
        """cd /Users/akifduran/Desktop/new_project/frontend-files
printf 'VITE_API_BASE_URL=http://127.0.0.1:8000/api\\n' > .env
npm install
npm run dev""",
    )
    add_text(
        doc,
        "API dokümantasyonu backend çalışırken `http://127.0.0.1:8000/docs` adresindeki Swagger UI üzerinden incelenebilir.",
    )

    doc.add_heading("16. Sunum İçin Hazır Anlatım", level=1)
    doc.add_heading("16.1 90 Saniyelik Teknik Konuşma", level=2)
    add_callout(
        doc,
        "Konuşma metni",
        "Backendimizi FastAPI ve PostgreSQL ile geliştirdik. Kullanıcı, Güz veya Bahar dönemini seçip ortak ders dosyasını yüklediğinde sistem tek komutla iki aşamalı çalışıyor. "
        "İlk aşamada ortak dersler OR-Tools CP-SAT ile planlanıyor. Bu aşamada oluşan öğretim üyesi dolulukları sabitleniyor. İkinci aşamada 12 mühendislik bölümünün programı sırayla oluşturuluyor ve her çözüm sonraki bölümlere çakışma bilgisi aktarıyor. "
        "Takvim, sabit/yasak saatler, bağlantılı dersler ve split tercihleri PostgreSQL'de tutuluyor; dolayısıyla algoritma tarayıcıdaki geçici state'e değil kayıtlı kurallara göre çalışıyor. "
        "Çözümden sonra bağımsız validator saat, ardışıklık, öğretim üyesi ve dönem-şube kurallarını yeniden kontrol ediyor. Sonuçlar ortak ve bölüm bazında JSONB olarak saklanıp frontend'e gönderiliyor.",
        fill="F4F6F9",
    )
    doc.add_heading("16.2 Olası Jüri Soruları", level=2)
    qa_rows = [
        ["Neden iki aşama?", "Ortak dersler tüm bölümleri etkilediği için önce tek kez yerleştirilir; dolulukları bölüm çözümlerine aktarılır."],
        ["Neden OR-Tools?", "Problem çok sayıda ayrık karar ve çakışmama kuralı içerir; CP-SAT uygulanabilir çözümü sistematik arar."],
        ["Kurallar nerede?", "Takvim, rule switch, zaman kısıtı, split ve linked group değerleri PostgreSQL'dedir."],
        ["Çözümün doğru olduğunu nasıl biliyorsunuz?", "CP-SAT modelinden bağımsız validator H1-H8 kontrollerini somut çıktı üzerinde tekrar yapar."],
        ["Excel yüklemek neden gerekli?", "Ortak ders satırları istekte gelir; bölüm satırları mevcut sürümde seçilen dönem klasöründeki sunucu Excel'lerinden okunur."],
        ["TBA nedir?", "Ders saati oluşmuş ancak fiziksel sınıf atama bileşeni bu akışa bağlanmamış demektir."],
        ["Programlar nerede saklanıyor?", "`published_schedule_documents.courses` JSONB alanında dönem ve bölüm bazında."],
        ["Aynı hoca iki bölümde çakışır mı?", "Bölümler sırayla çözülürken hoca doluluğu aktarılır; ayrıca birleşik sonuçta global kontrol yapılır."],
    ]
    add_table(doc, ["Soru", "Kısa cevap"], qa_rows, [2.25, 4.25], 8.8)

    doc.add_heading("17. Sonuç", level=1)
    add_text(
        doc,
        "OptiSched backend; HTTP servisleri, PostgreSQL tabanlı kural yönetimi ve CP-SAT çizelgeleme çekirdeğini çalışan bir bütün halinde birleştirir. "
        "Mimarinin güçlü tarafı, kısıtların kalıcı olması, ortak/bölüm programlarının tek çalıştırmada üretilmesi ve sonucun bağımsız doğrulamadan geçirilmesidir. "
        "Üretim seviyesine geçişte öncelikli konular sınıf atama entegrasyonu, bütün kritik endpointlerde rol tabanlı yetkilendirme ve yüklenen bölüm dosyalarının backend veri akışına tam bağlanmasıdır.",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
